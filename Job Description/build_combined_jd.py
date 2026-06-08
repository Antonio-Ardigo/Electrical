"""Build a single combined Word document from the three job-description files.

Reads the per-role Markdown JDs and renders them into one .docx with a title page.
Run:  python3 build_combined_jd.py
Output: WSM-Job-Descriptions-Combined.docx
"""
import os
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

HERE = os.path.dirname(__file__)
FILES = [
    "project-manager-electrical.md",
    "electrical-engineer.md",
    "electrical-technician.md",
]
OUT = os.path.join(HERE, "WSM-Job-Descriptions-Combined.docx")

LINK = re.compile(r"\[([^\]]+)\]\([^)]+\)")
INLINE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`)")


def add_runs(p, text):
    text = LINK.sub(r"\1", text)
    for part in INLINE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            p.add_run(part[2:-2]).bold = True
        elif part.startswith("`") and part.endswith("`"):
            p.add_run(part[1:-1])
        else:
            p.add_run(part)


def is_sep(cells):
    return all(re.fullmatch(r":?-{2,}:?", c.strip() or "-") for c in cells) and any("-" in c for c in cells)


def split_row(line):
    s = line.strip()
    if s.startswith("|"):
        s = s[1:]
    if s.endswith("|"):
        s = s[:-1]
    return [c.strip() for c in s.split("|")]


def add_table(doc, rows):
    # rows: list of cell-lists; detect header via separator second row
    header = None
    body = rows
    if len(rows) >= 2 and is_sep(rows[1]):
        header = rows[0]
        body = rows[2:]
        if all(c.strip() == "" for c in header):
            header = None  # empty header -> skip
    ncol = max(len(r) for r in rows)
    t = doc.add_table(rows=0, cols=ncol)
    t.style = "Table Grid"
    if header:
        cells = t.add_row().cells
        for i in range(ncol):
            cells[i].paragraphs[0].text = ""
            run = cells[i].paragraphs[0].add_run(header[i] if i < len(header) else "")
            run.bold = True
    for r in body:
        cells = t.add_row().cells
        for i in range(ncol):
            para = cells[i].paragraphs[0]
            add_runs(para, r[i] if i < len(r) else "")
    doc.add_paragraph()


def convert(doc, md, first):
    lines = md.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        s = line.strip()
        if not s:
            i += 1
            continue
        if s == "---":
            i += 1
            continue
        # table block
        if s.startswith("|"):
            block = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                block.append(split_row(lines[i]))
                i += 1
            add_table(doc, block)
            continue
        # headings
        if s.startswith("# "):
            if not first:
                doc.add_page_break()
            first = False
            doc.add_heading(s[2:].strip(), level=1)
        elif s.startswith("## "):
            doc.add_heading(s[3:].strip(), level=2)
        elif s.startswith("### "):
            doc.add_heading(s[4:].strip(), level=3)
        elif s.startswith("> "):
            p = doc.add_paragraph(style="Intense Quote")
            add_runs(p, s[2:].strip())
        elif s.startswith("- ") or s.startswith("* "):
            p = doc.add_paragraph(style="List Bullet")
            add_runs(p, s[2:].strip())
        else:
            p = doc.add_paragraph()
            add_runs(p, s)
        i += 1
    return first


doc = Document()
doc.styles["Normal"].font.name = "Calibri"
doc.styles["Normal"].font.size = Pt(10.5)

# ---- title page ----
t = doc.add_heading("Job Descriptions", level=0)
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.LEFT
r = sub.add_run("Electrical Single Point of Failure (SPOF) Assessment — "
                "21 water / desalination plants, Central & Eastern sector — 4-month campaign")
r.italic = True
doc.add_paragraph()
add_runs(doc.add_paragraph(),
         "**Team:** 1 Project Manager (Electrical) + 2 Electrical Engineers; "
         "optional Electrical Technician (field support).")
add_runs(doc.add_paragraph(),
         "**Engineers — required:** a complete plant calculation package "
         "(load-flow, short-circuit, protection coordination, earthing & lightning). "
         "**Preferable (not required):** arc-flash / incident-energy (IEEE 1584); "
         "RAM / RBD availability analysis.")
add_runs(doc.add_paragraph(),
         "**Key deliverables — the forms:** for each plant, a document inventory, a "
         "failure-mode / criticality (FMECA) form, a recommendations register, and a "
         "site-visit report.")
add_runs(doc.add_paragraph(),
         "**Common requirements:** LV & MV; plant assessment (SPOF / reliability / "
         "condition); SCE accreditation (to be discussed); Arabic & English; QHSE and "
         "host permit-to-work; non-intrusive surveys; NDA and national cybersecurity "
         "data handling.")

first = True
for fn in FILES:
    md = open(os.path.join(HERE, fn), encoding="utf-8").read()
    first = convert(doc, md, first)

doc.save(OUT)
print("wrote", OUT)
